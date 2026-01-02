"""
Basic DOCX Document Reader

This module provides basic functionality for loading and reading DOCX documents,
with file validation and error handling capabilities.
"""

import os
import logging
import re
from typing import List, Dict, Any, Optional
from pathlib import Path
from docx import Document
from docx.document import Document as DocxDocument
from docx.table import Table
from docx.text.paragraph import Paragraph

from faq.rag.interfaces.base import DocumentStructure, ValidationResult

logger = logging.getLogger(__name__)


class DOCXDocumentReader:
    """Basic DOCX document reader with validation and error handling."""
    
    def __init__(self):
        """Initialize the document reader."""
        self.supported_extensions = ['.docx']
        self.max_file_size_mb = 50  # Maximum file size in MB
    
    def validate_file(self, file_path: str) -> ValidationResult:
        """
        Validate DOCX file before processing.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            ValidationResult with validation status and any errors
        """
        errors = []
        warnings = []
        
        try:
            # Check if file exists
            if not os.path.exists(file_path):
                errors.append(f"File does not exist: {file_path}")
                return ValidationResult(
                    is_valid=False,
                    errors=errors,
                    warnings=warnings,
                    metadata={"file_path": file_path}
                )
            
            # Check file extension
            file_extension = Path(file_path).suffix.lower()
            if file_extension not in self.supported_extensions:
                errors.append(f"Unsupported file extension: {file_extension}. Supported: {self.supported_extensions}")
            
            # Check file size
            file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
            if file_size_mb > self.max_file_size_mb:
                errors.append(f"File too large: {file_size_mb:.2f}MB. Maximum allowed: {self.max_file_size_mb}MB")
            elif file_size_mb > 10:  # Warning for large files
                warnings.append(f"Large file detected: {file_size_mb:.2f}MB. Processing may take longer.")
            
            # Check if file is readable
            try:
                with open(file_path, 'rb') as f:
                    # Try to read first few bytes to ensure file is accessible
                    f.read(1024)
            except PermissionError:
                errors.append(f"Permission denied: Cannot read file {file_path}")
            except Exception as e:
                errors.append(f"File access error: {str(e)}")
            
            # Try to open as DOCX to validate format
            if not errors:
                try:
                    Document(file_path)
                except Exception as e:
                    errors.append(f"Invalid DOCX format: {str(e)}")
            
            return ValidationResult(
                is_valid=len(errors) == 0,
                errors=errors,
                warnings=warnings,
                metadata={
                    "file_path": file_path,
                    "file_size_mb": file_size_mb,
                    "file_extension": file_extension
                }
            )
            
        except Exception as e:
            logger.error(f"Unexpected error during file validation: {str(e)}")
            return ValidationResult(
                is_valid=False,
                errors=[f"Validation error: {str(e)}"],
                warnings=warnings,
                metadata={"file_path": file_path}
            )
    
    def load_document(self, file_path: str) -> Optional[DocxDocument]:
        """
        Load a DOCX document with error handling.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Document object if successful, None if failed
        """
        try:
            # Validate file first
            validation_result = self.validate_file(file_path)
            if not validation_result.is_valid:
                logger.error(f"File validation failed: {validation_result.errors}")
                return None
            
            # Log warnings if any
            for warning in validation_result.warnings:
                logger.warning(warning)
            
            # Load the document
            logger.info(f"Loading DOCX document: {file_path}")
            document = Document(file_path)
            logger.info(f"Successfully loaded document with {len(document.paragraphs)} paragraphs and {len(document.tables)} tables")
            
            return document
            
        except Exception as e:
            logger.error(f"Failed to load DOCX document {file_path}: {str(e)}")
            return None
    
    def extract_basic_text(self, document: DocxDocument) -> List[str]:
        """
        Extract basic text content from document paragraphs.
        
        Args:
            document: Loaded DOCX document
            
        Returns:
            List of paragraph texts
        """
        try:
            paragraphs = []
            for paragraph in document.paragraphs:
                text = paragraph.text.strip()
                if text:  # Only include non-empty paragraphs
                    paragraphs.append(text)
            
            logger.info(f"Extracted {len(paragraphs)} non-empty paragraphs")
            return paragraphs
            
        except Exception as e:
            logger.error(f"Failed to extract basic text: {str(e)}")
            return []
    
    def extract_table_data(self, document: DocxDocument) -> List[Dict[str, Any]]:
        """
        Extract data from all tables in the document.
        
        Args:
            document: Loaded DOCX document
            
        Returns:
            List of table data dictionaries
        """
        try:
            tables_data = []
            
            for table_idx, table in enumerate(document.tables):
                table_data = {
                    "table_index": table_idx,
                    "rows": [],
                    "row_count": len(table.rows),
                    "col_count": len(table.columns) if table.rows else 0
                }
                
                for row_idx, row in enumerate(table.rows):
                    row_data = {
                        "row_index": row_idx,
                        "cells": []
                    }
                    
                    for cell_idx, cell in enumerate(row.cells):
                        cell_text = cell.text.strip()
                        row_data["cells"].append({
                            "cell_index": cell_idx,
                            "text": cell_text
                        })
                    
                    table_data["rows"].append(row_data)
                
                tables_data.append(table_data)
            
            logger.info(f"Extracted data from {len(tables_data)} tables")
            return tables_data
            
        except Exception as e:
            logger.error(f"Failed to extract table data: {str(e)}")
            return []
    
    def analyze_document_structure(self, document: DocxDocument) -> DocumentStructure:
        """
        Analyze the overall structure of the document.
        
        Args:
            document: Loaded DOCX document
            
        Returns:
            DocumentStructure with analysis results
        """
        try:
            # Extract paragraphs
            paragraphs = self.extract_basic_text(document)
            
            # Extract tables
            tables = self.extract_table_data(document)
            
            # Analyze sections (basic implementation - could be enhanced)
            sections = []
            current_section = {"title": "Main Content", "paragraphs": []}
            
            for paragraph in paragraphs:
                # Improved heuristic: recognize 📌 and short bold/uppercase headers
                is_header = False
                if paragraph.startswith('📌'):
                    is_header = True
                elif len(paragraph) < 80 and (paragraph.endswith(':') or paragraph.isupper()):
                    is_header = True
                elif len(paragraph) < 100 and re.match(r'^\d+[️⃣⃣\.\)]*\s+', paragraph): # emoji numbers or standard numbers
                    # Only treat as header if it's not a question-like structure
                    if not paragraph.strip().endswith('?'):
                        is_header = True
                
                if is_header:
                    if current_section["paragraphs"]:
                        sections.append(current_section)
                    current_section = {"title": paragraph, "paragraphs": []}
                else:
                    current_section["paragraphs"].append(paragraph)
            
            if current_section["paragraphs"]:
                sections.append(current_section)
            
            # Identify lists (basic implementation)
            lists = []
            for paragraph in paragraphs:
                if (paragraph.startswith(('•', '-', '*', '1.', '2.', '3.', '4.', '5.')) or 
                    any(paragraph.startswith(f'{i}.') for i in range(1, 20))):
                    lists.append({
                        "type": "bullet" if paragraph.startswith(('•', '-', '*')) else "numbered",
                        "text": paragraph
                    })
            
            structure = DocumentStructure(
                document_type="docx",
                sections=sections,
                tables=tables,
                lists=lists,
                paragraphs=paragraphs
            )
            
            logger.info(f"Document structure analysis complete: {len(sections)} sections, {len(tables)} tables, {len(lists)} list items")
            return structure
            
        except Exception as e:
            logger.error(f"Failed to analyze document structure: {str(e)}")
            return DocumentStructure(
                document_type="docx",
                sections=[],
                tables=[],
                lists=[],
                paragraphs=[]
            )