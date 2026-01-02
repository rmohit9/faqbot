"""
Basic tests for DOCX Document Reader

These tests verify the core functionality of the document reader component.
"""

import os
import tempfile
from pathlib import Path
from docx import Document

from .document_reader import DOCXDocumentReader


def test_file_validation():
    """Test file validation functionality."""
    reader = DOCXDocumentReader()
    
    # Test non-existent file
    result = reader.validate_file("nonexistent.docx")
    assert not result.is_valid
    assert "File does not exist" in result.errors[0]
    
    # Test invalid extension
    with tempfile.NamedTemporaryFile(suffix=".txt", delete=False) as tmp:
        tmp.write(b"test content")
        tmp_path = tmp.name
    
    try:
        result = reader.validate_file(tmp_path)
        assert not result.is_valid
        assert "Unsupported file extension" in result.errors[0]
    finally:
        os.unlink(tmp_path)


def test_basic_document_loading():
    """Test basic document loading and text extraction."""
    reader = DOCXDocumentReader()
    
    # Create a simple test document
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = tmp.name
    
    try:
        # Create a simple DOCX document
        doc = Document()
        doc.add_paragraph("This is a test paragraph.")
        doc.add_paragraph("This is another test paragraph.")
        
        # Add a simple table
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Question"
        table.cell(0, 1).text = "Answer"
        table.cell(1, 0).text = "What is this?"
        table.cell(1, 1).text = "This is a test."
        
        doc.save(tmp_path)
        
        # Test loading
        loaded_doc = reader.load_document(tmp_path)
        assert loaded_doc is not None
        
        # Test text extraction
        paragraphs = reader.extract_basic_text(loaded_doc)
        assert len(paragraphs) >= 2
        assert "This is a test paragraph." in paragraphs
        
        # Test table extraction
        tables = reader.extract_table_data(loaded_doc)
        assert len(tables) == 1
        assert tables[0]["row_count"] == 2
        assert tables[0]["col_count"] == 2
        
        # Test structure analysis
        structure = reader.analyze_document_structure(loaded_doc)
        assert structure.document_type == "docx"
        assert len(structure.paragraphs) >= 2
        assert len(structure.tables) == 1
        
    finally:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)


if __name__ == "__main__":
    test_file_validation()
    test_basic_document_loading()
    print("All tests passed!")