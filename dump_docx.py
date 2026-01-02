from docx import Document
import sys
import os

def dump_docx(path):
    if not os.path.exists(path):
        print(f"File not found: {path}")
        return
    
    doc = Document(path)
    print(f"Document: {path}")
    print(f"Number of paragraphs: {len(doc.paragraphs)}")
    print("-" * 30)
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text:
            print(f"P{i}: {text}")
    
    print("-" * 30)
    print(f"Number of tables: {len(doc.tables)}")
    for i, table in enumerate(doc.tables):
        print(f"Table {i}: {len(table.rows)} rows, {len(table.columns)} columns")
        for row in table.rows:
            print([cell.text.strip() for cell in row.cells])

if __name__ == "__main__":
    if len(sys.argv) > 1:
        dump_docx(sys.argv[1])
    else:
        print("Usage: python dump_docx.py <docx_path>")
